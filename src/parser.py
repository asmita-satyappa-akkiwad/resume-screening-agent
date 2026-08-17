"""
parser.py
---------
Turns a resume FILE (pdf/docx/txt) into plain TEXT.

This is step 1 of the pipeline: "Agent detects file type -> Agent
extracts resume text". Everything after this file works purely with
strings, so parser.py is the only place that needs to know about
file formats.
"""

import os
from src.utils import read_text_file, clean_text

# fitz is the import name for the PyMuPDF library
import fitz
from docx import Document


class ParsingError(Exception):
    """Raised when a resume file can't be read at all."""


def extract_text(file_path: str) -> str:
    """
    Reads a resume file and returns its plain text content,
    regardless of whether it's a PDF, DOCX, or TXT file.

    This is the single function the rest of the app calls -
    it hides the "how" (which library, which format) from
    everything downstream.
    """
    if not os.path.exists(file_path):
        raise ParsingError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            text = _extract_from_pdf(file_path)
        elif ext == ".docx":
            text = _extract_from_docx(file_path)
        elif ext == ".txt":
            text = read_text_file(file_path)
        else:
            raise ParsingError(
                f"Unsupported file type '{ext}'. Only .pdf, .docx, .txt are supported."
            )
    except ParsingError:
        raise
    except Exception as e:
        # Wrap any library-specific error in our own exception type so
        # main.py can catch ONE thing instead of five different
        # library exceptions. This is what "handle errors gracefully"
        # means in practice.
        raise ParsingError(f"Failed to parse {file_path}: {e}")

    text = clean_text(text)

    if not text or len(text.strip()) < 20:
        raise ParsingError(
            f"{file_path} produced little/no text — file may be empty, "
            f"image-only (scanned), or corrupted."
        )

    return text


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def _extract_from_docx(file_path: str) -> str:
    """Extract text from a .docx using python-docx.

    We pull text from both normal paragraphs AND tables, because a
    lot of resume templates put skills/experience inside tables.
    """
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)
